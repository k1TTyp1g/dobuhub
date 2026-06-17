"""
DocuHub 核心排版引擎
======================
使用 python-docx 对 .docx 文件执行一键式规范化排版。
核心原则：非破坏性——始终另存为新文件，绝不修改原稿。

技术选型说明：
  - python-docx 是目前对 .docx 支持最稳定、API 最完善的 Python 库
  - 对比 Java 的 Apache POI：python-docx 在简单文档批量处理上优势明显
  - 对比 docx4j：python-docx 学习成本低，功能满足需求
  - .doc（旧格式）需先转换为 .docx（见 doc_converter.py）
"""
import re
import os
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models.template import FormatTemplate

# ── 单位转换 ─────────────────────────────────────────────
# python-docx 使用 EMU（English Metric Unit）
# 1 cm = 360000 EMU,  1 pt = 12700 EMU
CM_TO_EMU = 360000
PT_TO_EMU = 12700


class StyleEngine:
    """文档样式引擎：将模板规则应用到 .docx 文件"""

    def __init__(self, template: FormatTemplate):
        self.template = template
        self.stats = {"ok": 0, "skip": 0, "error": 0, "details": []}

    # ── 公开接口 ─────────────────────────────────────────

    def format_document(self, src_path: str, dst_path: Optional[str] = None) -> str:
        """
        对单个 .docx 应用模板格式。

        参数：
            src_path: 源文件路径
            dst_path: 输出路径（None 时自动生成）

        返回：
            输出文件路径
        """
        if dst_path is None:
            src = Path(src_path)
            dst_path = str(src.parent / f"{src.stem}_已格式化{src.suffix}")

        try:
            doc = Document(src_path)
            self._apply_page_setup(doc)
            self._apply_body_style(doc)
            self._apply_headings(doc)
            doc.save(dst_path)
            self.stats["ok"] += 1
            self.stats["details"].append(f"✅ {Path(src_path).name} → {Path(dst_path).name}")
        except Exception as e:
            self.stats["error"] += 1
            self.stats["details"].append(f"❌ {Path(src_path).name}: {e}")
            raise

        return dst_path

    def batch_format(self, file_list: list, output_dir: Optional[str] = None,
                     progress_callback=None) -> list:
        """
        批量格式化多个文件。

        参数：
            file_list: 源文件路径列表
            output_dir: 输出目录（None 时与源文件同目录）
            progress_callback: 进度回调函数 f(当前序号, 总数, 文件名)

        返回：
            输出文件路径列表
        """
        results = []
        self.stats = {"ok": 0, "skip": 0, "error": 0, "details": []}

        for i, src_path in enumerate(file_list):
            filename = Path(src_path).name

            if progress_callback:
                progress_callback(i + 1, len(file_list), filename)

            # 跳过非 .docx 文件
            ext = Path(src_path).suffix.lower()
            if ext not in (".docx",):
                self.stats["skip"] += 1
                self.stats["details"].append(f"⏭ {filename}: 非 .docx 格式，跳过")
                results.append(None)
                continue

            # 确定输出路径
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                src_name = Path(src_path).name
                dst_path = os.path.join(output_dir, src_name)
            else:
                dst_path = None

            try:
                out = self.format_document(src_path, dst_path)
                results.append(out)
            except Exception as e:
                results.append(None)

        return results

    # ── 内部实现 ─────────────────────────────────────────

    def _apply_page_setup(self, doc: Document):
        """应用页面设置（纸张大小、页边距）"""
        tpl = self.template
        section = doc.sections[0]

        # 纸张大小
        if tpl.page.paper_size.upper() == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)

        # 页边距（cm → EMU）
        section.top_margin = int(tpl.page.margin_top * CM_TO_EMU)
        section.bottom_margin = int(tpl.page.margin_bottom * CM_TO_EMU)
        section.left_margin = int(tpl.page.margin_left * CM_TO_EMU)
        section.right_margin = int(tpl.page.margin_right * CM_TO_EMU)

    def _apply_body_style(self, doc: Document):
        """将正文格式应用到所有非标题段落"""
        tpl = self.template
        level_keywords = [f"Heading {i}" for i in range(1, 10)]

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            # 跳过标题段落（交给 _apply_headings 处理）
            if any(kw in style_name for kw in level_keywords):
                continue

            # 跳过空段落（但保留）
            self._set_paragraph_format(para, tpl.body)

    def _apply_headings(self, doc: Document):
        """识别并格式化标题段落"""
        tpl = self.template

        # python-docx 默认使用 'Heading 1', 'Heading 2', 'Heading 3' 等样式名
        # WPS 和 Word 生成的文档都遵循这个命名规范
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            # 自动识别标题级别
            for level in [1, 2, 3]:
                if style_name == f"Heading {level}":
                    hs = tpl.headings.get(level, tpl.headings[1])
                    self._set_paragraph_format(para, hs)
                    break

    def _set_paragraph_format(self, para, style_obj):
        """
        将样式设置应用到单个段落。
        这是核心格式化函数，处理：对齐、缩进、行距、段间距、字体、字号、加粗。
        """
        pf = para.paragraph_format

        # ── 对齐方式 ──
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = align_map.get(style_obj.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

        # ── 首行缩进（2字符 = 2 * 字体大小 pt） ──
        if hasattr(style_obj, 'first_line_indent') and style_obj.first_line_indent > 0:
            indent_pt = style_obj.first_line_indent * style_obj.font_size
            pf.first_line_indent = Pt(indent_pt)
        else:
            pf.first_line_indent = Pt(0)

        # ── 行距 ──
        try:
            spacing_type = getattr(style_obj, 'line_spacing_type', 'fixed')
            spacing_val = getattr(style_obj, 'line_spacing', 20)

            if spacing_type == "fixed":
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(spacing_val)
            elif spacing_type == "multiple":
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = spacing_val
            elif spacing_type == "single":
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                pf.line_spacing = 1.0
            else:
                from docx.enum.text import WD_LINE_SPACING
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(spacing_val)
        except Exception:
            pass

        # ── 段前段后间距 ──
        pf.space_before = Pt(getattr(style_obj, 'space_before', 0))
        pf.space_after = Pt(getattr(style_obj, 'space_after', 0))

        # ── 遍历所有 run 设置字体 ──
        for run in para.runs:
            self._set_run_font(run, style_obj)

        # 如果段落没有 run（空段落但有文本？通过 xml 写字体）
        if not para.runs and para.text.strip():
            run = para.add_run(para.text)
            self._set_run_font(run, style_obj)

    def _set_run_font(self, run, style_obj):
        """设置 run 级别的字体属性"""
        font = run.font

        # 字号
        font.size = Pt(style_obj.font_size)

        # 加粗 / 斜体
        font.bold = style_obj.bold
        font.italic = style_obj.italic

        # 中西文字体分开设置（关键：确保 WPS 和 Word 都正确显示）
        try:
            # 设置西文字体
            font.name = style_obj.font_name_en

            # 设置中文字体（通过 XML 操作，python-docx 没有直接 API）
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
                rpr.insert(0, rfonts)

            rfonts.set(qn('w:eastAsia'), style_obj.font_name_cn)
            rfonts.set(qn('w:ascii'), style_obj.font_name_en)
            rfonts.set(qn('w:hAnsi'), style_obj.font_name_en)
        except Exception:
            # 如果 XML 操作失败，至少设置了 font.name
            pass
